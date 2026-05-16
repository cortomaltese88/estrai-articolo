#!/usr/bin/env python3
"""
server.py — backend locale per estrai-articolo
"""

import http.server
import json
import re
import os
import io
import urllib.request
import urllib.error

PORT = 7432
BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def estrai(url):
    headers = {
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 Chrome/120.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Accept-Language': 'it-IT,it;q=0.9,en;q=0.8',
    }
    req = urllib.request.Request(url, headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=15) as r:
            html = r.read().decode('utf-8', errors='replace')
    except Exception as e:
        return {'errore': f'Errore download: {e}'}

    matches = re.findall(r'application/ld\+json[^>]*>(.*?)</script>', html, re.DOTALL)
    articolo = None
    for m in matches:
        try:
            d = json.loads(m)
            if 'articleBody' in d:
                articolo = d
                break
        except:
            pass

    testo = ''
    titolo = ''
    data = ''
    autore = ''
    testata = ''
    url_can = ''

    if articolo:
        titolo  = articolo.get('headline', '')
        data    = articolo.get('datePublished', '')[:10]
        if isinstance(articolo.get('author'), dict):
            autore = articolo['author'].get('name', '')
        if isinstance(articolo.get('publisher'), dict):
            testata = articolo['publisher'].get('name', '')
        url_can = articolo.get('url', url)
        testo = articolo.get('articleBody', '')
    else:
        m = re.search(r'"articleBody"\s*:\s*"(.*?)(?<!\\)"(?:\s*,|\s*\})', html, re.DOTALL)
        if m: testo = m.group(1)
        t = re.search(r'"headline"\s*:\s*"(.*?)"', html)
        if t: titolo = t.group(1)
        d = re.search(r'"datePublished"\s*:\s*"(.*?)"', html)
        if d: data = d.group(1)[:10]
        a = re.search(r'"author".*?"name"\s*:\s*"(.*?)"', html, re.DOTALL)
        if a: autore = a.group(1)
        p = re.search(r'"publisher".*?"name"\s*:\s*"(.*?)"', html, re.DOTALL)
        if p: testata = p.group(1)
        u = re.search(r'<link rel="canonical" href="(https://[^"]+)"', html)
        if u: url_can = u.group(1)

    if not testo:
        return {'errore': 'Contenuto non estraibile dalla pagina fornita.'}

    testo = testo.replace('&nbsp;', ' ').replace('\\n', '\n').replace('\\"', '"')
    testo = re.sub(r'<[^>]+>', '', testo)
    testo = re.sub(r'[ \t]+', ' ', testo)
    testo = re.sub(r'\n{3,}', '\n\n', testo).strip()

    return {
        'titolo': titolo, 'data': data, 'autore': autore,
        'testata': testata, 'url': url_can or url,
        'testo': testo, 'errore': ''
    }


def genera_docx(data):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margini pagina
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Intestazione rassegna stampa
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run('RASSEGNA STAMPA')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Metadati
    meta = [
        ('Testata', data.get('testata', '')),
        ('Data',    data.get('data', '')),
        ('Autore',  data.get('autore', '')),
        ('URL',     data.get('url', '')),
    ]
    for k, v in meta:
        if v:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            lb = p.add_run(f'{k}:  ')
            lb.bold = True
            lb.font.size = Pt(10)
            vr = p.add_run(v)
            vr.font.size = Pt(10)

    # Separatore
    doc.add_paragraph('─' * 60)

    # Titolo
    titolo_p = doc.add_paragraph()
    titolo_p.paragraph_format.space_before = Pt(8)
    titolo_p.paragraph_format.space_after  = Pt(12)
    tr = titolo_p.add_run(data.get('titolo', ''))
    tr.bold = True
    tr.font.size = Pt(14)

    # Testo — paragrafi
    testo = data.get('testo', '')
    paragrafi = testo.split('\n')
    paragrafi = [p.strip() for p in paragrafi if p.strip()]
    if len(paragrafi) <= 1:
        paragrafi = re.sub(r'([.!?])\s{2,}([A-ZAÀÈÉÌÒÙ])', r'\1\n\2', testo).split('\n')
        paragrafi = [p.strip() for p in paragrafi if p.strip()]

    for par in paragrafi:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(par)
        r.font.size = Pt(11)

    # Separatore finale
    doc.add_paragraph('─' * 60)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class Handler(http.server.BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass

    def do_GET(self):
        if self.path == '/' or self.path == '/index.html':
            gui_path = os.path.join(BASE_DIR, 'gui.html')
            try:
                with open(gui_path, 'rb') as f:
                    content = f.read()
                self.send_response(200)
                self.send_header('Content-Type', 'text/html; charset=utf-8')
                self.send_header('Content-Length', len(content))
                self.end_headers()
                self.wfile.write(content)
            except FileNotFoundError:
                self.send_response(500)
                self.end_headers()
                self.wfile.write(b'ERRORE: gui.html non trovato accanto a server.py')
        else:
            self.send_response(404)
            self.end_headers()

    def do_POST(self):
        length = int(self.headers.get('Content-Length', 0))
        body   = self.rfile.read(length)
        data   = json.loads(body)

        if self.path == '/estrai':
            url = data.get('url', '').strip()
            result = estrai(url) if url else {'errore': 'URL mancante'}
            response = json.dumps(result, ensure_ascii=False).encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'application/json; charset=utf-8')
            self.send_header('Content-Length', len(response))
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(response)

        elif self.path == '/docx':
            try:
                docx_bytes = genera_docx(data)
                testata = data.get('testata', 'articolo').lower().replace(' ', '_')
                dt      = data.get('data', 'nd')
                slug    = re.sub(r'[^a-z0-9]', '_', data.get('titolo', '')[:40].lower())
                fname   = f"{testata}_{dt}_{slug}.docx"
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Disposition', f'attachment; filename="{fname}"')
                self.send_header('Content-Length', len(docx_bytes))
                self.send_header('Access-Control-Allow-Origin', '*')
                self.end_headers()
                self.wfile.write(docx_bytes)
            except Exception as e:
                err = json.dumps({'errore': str(e)}).encode()
                self.send_response(500)
                self.send_header('Content-Type', 'application/json')
                self.send_header('Content-Length', len(err))
                self.end_headers()
                self.wfile.write(err)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()


if __name__ == '__main__':
    print(f'Estrai Articolo — http://localhost:{PORT}')
    print(f'Directory: {BASE_DIR}')
    print('Ctrl+C per fermare.')
    httpd = http.server.HTTPServer(('localhost', PORT), Handler)
    httpd.serve_forever()
